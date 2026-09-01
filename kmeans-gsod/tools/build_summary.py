"""
Generate the one-page Executive Summary as a Microsoft Word document.

The assignment asks for the summary in Word. Writing it here rather than by hand
means the numbers come from one place and the document can be regenerated if the
analysis changes.

Not part of the application. Needs one extra package:

    pip install python-docx
    python tools\\build_summary.py

Length is tuned to fit a single page: 0.7 inch margins, 10 pt body, and about 420
words of prose plus one table. Verified at one page with Word's own pagination.
If you edit the text, keep an eye on that budget -- the first draft ran to 641
words and spilled onto a second page.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parents[1] / "EXECUTIVE_SUMMARY.docx"

BODY_SIZE = Pt(10)
ACCENT = RGBColor(0x1F, 0x3B, 0x73)

CLUSTER_TABLE = [
    ("Group", "Stations", "Mean temp", "Seasonality", "Rainfall", "Mean |latitude|"),
    ("Polar / tundra", "23", "-1.0 °C", "20.4 °C", "621 mm", "60.5°"),
    ("Subarctic continental", "81", "-0.5 °C", "39.7 °C", "570 mm", "56.9°"),
    ("Arid / semi-arid", "31", "12.4 °C", "26.9 °C", "211 mm", "35.6°"),
    ("Maritime temperate", "114", "16.5 °C", "15.9 °C", "1,011 mm", "36.2°"),
    ("Tropical humid", "24", "27.1 °C", "3.6 °C", "2,802 mm", "11.4°"),
]

SECTIONS = [
    ("Purpose",
     "Can a machine learning algorithm, given only a year of daily weather "
     "readings, rediscover the world's climate zones without being told where any "
     "station is? The technique generalises: grouping things that behave alike "
     "when nobody has labelled them in advance is the problem behind customer "
     "segmentation and fault detection."),

    ("Dataset",
     "NOAA's Global Surface Summary of the Day, from the Registry of Open Data on "
     "AWS (s3://noaa-gsod-pds), holds daily readings from over 12,000 weather "
     "stations worldwide. The application sampled 531 stations across all 90 "
     "international reporting regions and retained the 273 with sufficiently "
     "complete records for 2023. The data is public and free; no AWS account was "
     "required."),

    ("Approach",
     "Each station's year of readings was condensed into ten descriptive numbers: "
     "average temperature, the swing between summer and winter, rainfall, "
     "humidity, snow, wind, fog and thunder. K-Means then grouped the stations by "
     "similarity across those measures.\n"
     "The critical decision was to withhold location. Latitude, longitude and "
     "elevation were never given to the algorithm, but kept back as an "
     "independent test: if the groups match geography anyway, the algorithm found "
     "something genuinely present in the weather rather than being handed the "
     "answer."),
]

FINDINGS_LEAD = ("The algorithm settled on five climate groups, and each is "
                 "immediately recognisable:")

FINDINGS_TAIL = (
    "Cluster membership explains 63.8% of the variation in distance from the "
    "equator, and the model never saw a coordinate. The most typical stations "
    "confirm it: Russian Arctic outposts and an Antarctic base in the polar group, "
    "Siberia and Fairbanks in subarctic, three cities on the Gobi Desert margin in "
    "arid, and Colombo, Manila and Thiruvananthapuram in tropical."
)

CLOSING = [
    ("Challenges",
     "Hidden missing data was the main obstacle, and it produced convincing but "
     "wrong results twice. NOAA does not leave gaps blank: a missing temperature "
     "is written as 9999.9, which turned one station's annual average into several "
     "hundred degrees. Worse, when a station reports no rainfall data at all, the "
     "file records 0.00 and flags it only in a separate column. That affected a "
     "fifth of all station-days and reported zero annual rainfall for a Scottish "
     "Highland pass receiving over 1,000 mm. Both faults are now covered by "
     "automated tests.\n"
     "Rainfall also overwhelmed the other measures, grouping a Norwegian station "
     "with the tropics because both are wet; a logarithmic scale restored balance. "
     "The standard method for choosing the number of groups proved unreliable too, "
     "its best answer beating the runner-up by 0.0007, which is noise rather than "
     "evidence. The count came from the point of diminishing returns instead, and "
     "three independent measures then agreed on five."),

    ("Conclusion",
     "Weather behaviour alone reconstructs the broad shape of world geography. Two "
     "caveats matter. Climate is a continuum, not a set of separate boxes: 42% of "
     "stations fall into the broad temperate group, and 4% sit near enough to a "
     "boundary that their assignment could reasonably go either way. Two stations "
     "at the same Montreal airport, reporting within 0.2 °C of each other, ended "
     "up in different groups. These results also describe 2023 only, not a "
     "long-term climate normal."),
]

SOURCE = ("Source: NOAA National Centers for Environmental Information. Global "
          "Surface Summary of the Day. Registry of Open Data on AWS. "
          "https://registry.opendata.aws/noaa-gsod/ (accessed 31 August 2026).")


def configure_document() -> Document:
    """Narrow margins and a compact body font, so the content fits one page."""
    document = Document()

    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.7)
    section.left_margin = section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = BODY_SIZE
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.0

    return document


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = ACCENT


def add_body(document: Document, text: str) -> None:
    """Blank-line-separated text becomes separate paragraphs."""
    for chunk in text.split("\n"):
        document.add_paragraph(chunk.strip())


def add_cluster_table(document: Document) -> None:
    table = document.add_table(rows=len(CLUSTER_TABLE), cols=len(CLUSTER_TABLE[0]))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_index, row in enumerate(CLUSTER_TABLE):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(1)
            if col_index > 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run(value)
            run.font.size = Pt(9.5)
            run.bold = row_index == 0

    widths = [Inches(1.55), Inches(0.7), Inches(0.85), Inches(0.95),
              Inches(0.85), Inches(1.15)]
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width


def build() -> Path:
    document = configure_document()

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(1)
    title_run = title.add_run("Discovering the World's Climate Zones from Weather Data Alone")
    title_run.bold = True
    title_run.font.size = Pt(15)
    title_run.font.color.rgb = ACCENT

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle_run = subtitle.add_run(
        "Executive Summary  |  Module 4 Assignment: K-Means Python Application"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)

    for heading, text in SECTIONS:
        add_heading(document, heading)
        add_body(document, text)

    add_heading(document, "Findings")
    add_body(document, FINDINGS_LEAD)
    add_cluster_table(document)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    add_body(document, FINDINGS_TAIL)

    for heading, text in CLOSING:
        add_heading(document, heading)
        add_body(document, text)

    source = document.add_paragraph()
    source.paragraph_format.space_before = Pt(6)
    source_run = source.add_run(SOURCE)
    source_run.italic = True
    source_run.font.size = Pt(8.5)

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path} ({path.stat().st_size:,} bytes)")
